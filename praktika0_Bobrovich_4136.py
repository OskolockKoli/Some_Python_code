import os
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import Font

# Функция для чтения данных из файла info.txt
def read_data_from_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        data = file.read().splitlines()
    return data

# Функция для получения содержимого лабораторной работы
def get_lab_content(lab_number):
    try:
        with open(f"{lab_number}.txt", 'r', encoding='utf-8') as lab_file:
            lab_content = lab_file.read().split('\n\n')
            lab_name = lab_content[0]
            lab_sections = lab_content[1:]
        return lab_name, lab_sections
    except FileNotFoundError:
        return None, None

# Функция для создания docx файлов
def create_docx_files(data):
    for line in data:
        student_records = line.split(';')
        for student_record in student_records:
            fields = student_record.split('/')
            full_name = fields[0].split()
            group_number = fields[1]
            lab_number = fields[2]
            current_date = datetime.datetime.now().strftime('%d.%m.%Y')
            current_year = datetime.datetime.now().strftime('%Y')

            file_name = f"{' '.join(full_name)}_{group_number}_{lab_number}.docx"

            document = Document()
            lab_name, lab_sections = get_lab_content(lab_number)
            
            # Добавление текста с определенным шрифтом, размером и ориентацией
            title_paragraph = document.add_paragraph()
            title_text1 = f"МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ федеральное государственное автономное образовательное учреждение высшего образования «САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ АЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ»\n\nКАФЕДРА №  43\n\n"
            title_run1 = title_paragraph.add_run(title_text1)
            title_run1.font.name = 'Times New Roman'
            title_run1.font.size = Pt(12)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            title_paragraph = document.add_paragraph()
            title_text2 = f"ОТЧЁТ\nЗАЩИЩЁН С ОЦЕНКОЙ\nПРЕПОДАВАТЕЛЬ"
            title_run2 = title_paragraph.add_run(title_text2)
            title_run2.font.name = 'Times New Roman'
            title_run2.font.size = Pt(12)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            title_paragraph = document.add_paragraph()
            title_text3 = f"Старший преподаватель                                          М.Д. Поляк\nдолжность, уч. Степень, звание        подпись, дата    инициалы, фамилия\n\n\n\n"
            title_run3 = title_paragraph.add_run(title_text3)
            title_run3.font.name = 'Times New Roman'
            title_run3.font.size = Pt(12)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            title_paragraph = document.add_paragraph()
            title_text0 = f"Отчёт о лабораторной работе №{lab_number}.\n\n{lab_name}\nпо курсу: Операционные системы\n\n\n\n"
            title_run0 = title_paragraph.add_run(title_text0)
            title_run0.font.name = 'Times New Roman'
            title_run0.font.size = Pt(16)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            title_paragraph = document.add_paragraph()
            title_text4 = f"РАБОТУ ВЫПОЛНИЛ\nСТУДЕНТ ГР.      {group_number}                      {current_date}                               {' '.join(full_name)}\n                                                           подпись, дата                         инициалы, фамилия\n\n\n"
            title_run4 = title_paragraph.add_run(title_text4)
            title_run4.font.name = 'Times New Roman'
            title_run4.font.size = Pt(12)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            title_paragraph = document.add_paragraph()
            title_text5 = f"Санкт-Петербург {current_year}\n"
            title_run5 = title_paragraph.add_run(title_text5)
            title_run5.font.name = 'Times New Roman'
            title_run5.font.size = Pt(12)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if lab_name and lab_sections:
                # Добавление содержимого лабораторной работы
                for section in lab_sections:
                    section_paragraph = document.add_paragraph(section)
                    section_paragraph.style.font.name = 'Times New Roman'
                    section_paragraph.style.font.size = Pt(14)

            document.save(file_name)
            print(f"Создан файл отчёта {file_name}")

# Основная функция
def main():
    data = read_data_from_file('info.txt')
    create_docx_files(data)

if __name__ == "__main__":
    main()
