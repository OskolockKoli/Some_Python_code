import tkinter as tk
import pikepdf
import zlib
import datetime
import re
import sys
import hashlib
import docx
import os
import glob
import config
import keyboard
import pathlib
import fnmatch
import pikepdf
import binascii
import olefile
import openpyxl

from pkg_resources import DistributionNotFound
from pkg_resources import get_distribution as _get_distribution
from tkinter import messagebox as mb
from tkinter import filedialog
from tkinter import *
from PyPDF2 import PdfFileReader
from dateutil.tz import tzutc, tzoffset
from glob import glob
from docx import Document
from os import path
from config import *
from openpyxl import load_workbook

sys.path.append('../../')

pdf_date_pattern = re.compile(''.join([
    r"(D:)?",
    r"(?P<year>\d\d\d\d)",
    r"(?P<month>\d\d)",
    r"(?P<day>\d\d)",
    r"(?P<hour>\d\d)",
    r"(?P<minute>\d\d)",
    r"(?P<second>\d\d)",]))

def transform_date(date_str):
    """
    Преобразование даты в формате PDF, например "D: 20120321183444 + 07'00 '", в 
    удобное для использования дату и время.
    (Д: ГГГГММДДЧЧммССОХЧ'мм ')
    """
    global pdf_date_pattern
    match = re.match(pdf_date_pattern, date_str)
    if match:
        date_info = match.groupdict()
        for k, v in date_info.items():
            if v is None:
                pass
            elif k == 'tz_offset':
                date_info[k] = v.lower()
            else:
                date_info[k] = int(v)
        return datetime.datetime(**date_info)

def clicked():
    if selected.get() == 1:
        global folder_path
        #print(folder_path.get())
        folder_path.get()
        paths = []
        folder = folder_path.get()
        #print(folder)
        
        i = 0
        m = 0
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                i = i + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                m = m + 1
        
        j = 0
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                j = j + 1
        k = 0
        n = 0
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                k = k + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                n = n + 1
        
        doct = Document()
        d = i + j + k + m + n
        #print(d)
        table = doct.add_table(rows = d + 1, cols = 6)
        table.style = 'Table Grid'

        #создание ячеек
        for col in range(6):
            for row in range (d + 1):
                cell = table.cell(row, col)
                cell.text = str(0)

        #нумерация документов с первого встреченного
        for row in range(1, d + 1):
            col = 0
            cell = table.cell(row, col)
            cell.text = str(row)

        #заполнение названий документов (1)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1

        #заполнение последних дат изменений (2)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                currentFile = docx.Document(currentFile)
                properties = currentFile.core_properties
                col = 2
                cell = table.cell(row, col)
                cell.text = str(properties.modified)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 2
                cel = table.cell(row, col)
                cell.text = str(0)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                pdf = pikepdf.Pdf.open(currentFile)
                docinfo = pdf.docinfo
                col = 2
                cell = table.cell(row, col)
                o = 0
                for key, value in docinfo.items():
                    if str(value).startswith("D:"):
                        value = transform_date(str(pdf.docinfo["/CreationDate"]))
                        if o == 0:
                            if value != '':
                                #print(str(value))
                                cell.text = str(value)
                            else:
                                cell.text = str('-')
                        if o == 2:
                            if value != '':
                                #print(str(value))
                                cell.text = str(value)
                            else:
                                cell.text = str('-')
                    o = o + 1
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                col = 2
                cell = table.cell(row, col)
                i = currentFile.resolve().stem + ".xlsx"
                #print(i)
                file = load_workbook(currentFile)
                #print(file.properties.modified)
                j = file.properties.modified
                #print(j)
                cell.text = str(j)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                col = 2
                cell = table.cell(row, col)
                i = currentFile.resolve().stem + ".xls"
                #print(i)
                assert olefile.isOleFile(currentFile)
                ole = olefile.OleFileIO(currentFile)
                meta = ole.get_metadata()
                #print('Дата последнего сохранения:  '+str(meta.last_saved_time))
                cell.text = str(meta.last_saved_time)
                row = row + 1
                ole.close()
        
        #заполнение размера документов (3)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1

        #заполнение версий документов (4)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                i = docx.Document(i)
                properties = i.core_properties
                for col in range(5):
                    cell = table.cell(row, col)
                    if col == 4:
                        if properties.version != '':
                            cell.text = str(properties.version)
                        else:
                            cell.text = str('0')
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                #pdf = Pdf(i)
                #version = pdf.pdf_version
                cell = table.cell(row, col)
                cell.text = str(0)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                #pdf = Pdf(i)
                #version = pdf.pdf_version
                cell = table.cell(row, col)
                cell.text = str(0)
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                #pdf = Pdf(i)
                #version = pdf.pdf_version
                cell = table.cell(row, col)
                cell.text = str(0)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                #pdf = Pdf(i)
                #version = pdf.pdf_version
                cell = table.cell(row, col)
                cell.text = str(0)
                row = row + 1
        
        #заполнение кооннтрольной суммы документов (5)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 5
                cell = table.cell(row, col)
                file_name = i
                cell.text = str(hashlib.md5(open(file_name, 'rb').read()).hexdigest())
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 5
                cell = table.cell(row, col)
                file_name = i
                cell.text = str(hashlib.md5(open(file_name, 'rb').read()).hexdigest())
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 5
                cell = table.cell(row, col)
                file_name = i
                cell.text = str(hashlib.md5(open(file_name, 'rb').read()).hexdigest())
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 5
                cell = table.cell(row, col)
                file_name = i
                cell.text = str(hashlib.md5(open(file_name, 'rb').read()).hexdigest())
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 5
                cell = table.cell(row, col)
                file_name = i
                cell.text = str(hashlib.md5(open(file_name, 'rb').read()).hexdigest())
                row = row + 1

        #свойства документа
        for col in range (6):
            row = 0
            cell = table.cell(row, col)
            if col == 0:
                cell.text = str('Номер п/п')
            if col == 1:
                cell.text = str('Наименование электронного документа (файла)')
            if col == 2:
                cell.text = str('Дата и время последнего изменения документа (файла)')
            if col == 3:
                cell.text = str('Размер файла (байт)')
            if col == 4:
                cell.text = str('Версия')
            if col == 5:
                cell.text = str('MD5')

        doct.add_paragraph()

        table2 = doct.add_table(rows = 2, cols = 3)
        table2.style = 'Table Grid'

        #создание ячеек
        for col in range(3):
            for row in range (2):
                cell = table2.cell(row, col)
                cell.text = str('')
        
        #свойства документа
        for col in range (3):
            row = 0
            cell = table2.cell(row, col)
            if col == 1:
                cell.text = str('Исполнитель')
            if col == 2:
                cell.text = str('Подпись')

        doct.add_paragraph()

        table3 = doct.add_table(rows = 2, cols = 2)
        table3.style = 'Table Grid'

        #создание ячеек
        for col in range(2):
            for row in range (2):
                cell = table3.cell(row, col)
                cell.text = str('')
        
        #сохранение искомой таблицы
        doct.save('MD5 общая таблица.docx')
        
        lbl11.configure(text="Готово!1")
        lbl12.configure(text="Теперь готовый ИУЛ ждёт Вас в папке с программой")
    
    if selected.get() == 2:
        #print(folder_path.get())
        folder_path.get()
        paths = []
        folder = folder_path.get()
        #print(folder)
        i = 0
        m = 0
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                i = i + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                m = m + 1
        j = 0
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                j = j + 1
        k = 0
        n = 0
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                k = k + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                n = n + 1
        
        doct = Document()
        d = i + j + k + m + n
        #print(d)
        table = doct.add_table(rows = d + 1, cols = 6)
        table.style = 'Table Grid'

        #создание ячеек
        for col in range(6):
            for row in range (d + 1):
                cell = table.cell(row, col)
                cell.text = str(0)

        #нумерация документов с первого встреченного
        for row in range(1, d + 1):
            col = 0
            cell = table.cell(row, col)
            cell.text = str(row)

        #заполнение названий документов (1)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1

        #заполнение последних дат изменений (2)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                currentFile = docx.Document(currentFile)
                properties = currentFile.core_properties
                col = 2
                cell = table.cell(row, col)
                cell.text = str(properties.modified)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 2
                cell = table.cell(row, col)
                cell.text = str(0)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                pdf = pikepdf.Pdf.open(currentFile)
                docinfo = pdf.docinfo
                col = 2
                cell = table.cell(row, col)
                o = 0
                for key, value in docinfo.items():
                    if str(value).startswith("D:"):
                        value = transform_date(str(pdf.docinfo["/CreationDate"]))
                        if o == 0:
                            if value != '':
                                #print(str(value))
                                cell.text = str(value)
                            else:
                                cell.text = str('-')
                        if o == 2:
                            if value != '':
                                #print(str(value))
                                cell.text = str(value)
                            else:
                                cell.text = str('-')
                    o = o + 1
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                col = 2
                cell = table.cell(row, col)
                i = currentFile.resolve().stem + ".xlsx"
                #print(i)
                file = load_workbook(currentFile)
                #print(file.properties.modified)
                j = file.properties.modified
                #print(j)
                cell.text = str(j)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                col = 2
                cell = table.cell(row, col)
                i = currentFile.resolve().stem + ".xls"
                #print(i)
                assert olefile.isOleFile(currentFile)
                ole = olefile.OleFileIO(currentFile)
                meta = ole.get_metadata()
                #print('Дата последнего сохранения:  '+str(meta.last_saved_time))
                cell.text = str(meta.last_saved_time)
                row = row + 1
                ole.close()
        
        #заполнение размера документов (3)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1

        #заполнение версий документов (4)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                currentFile = docx.Document(currentFile)
                properties = currentFile.core_properties
                for col in range(5):
                    cell = table.cell(row, col)
                    if col == 4:
                        if properties.version != '':
                            cell.text = str(properties.version)
                        else:
                            cell.text = str('0')
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1
        
        #заполнение кооннтрольной суммы документов (5)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 5
                cell = table.cell(row, col)
                with open (i, 'rb') as fp:
                    file_name = fp.read()
                cell.text = str(hex(zlib.crc32(file_name)))
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 5
                cell = table.cell(row, col)
                with open (i, 'rb') as fp:
                    file_name = fp.read()
                cell.text = str(hex(zlib.crc32(file_name)))
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 5
                cell = table.cell(row, col)
                with open (i, 'rb') as fp:
                    file_name = fp.read()
                cell.text = str(hex(zlib.crc32(file_name)))
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 5
                cell = table.cell(row, col)
                with open (i, 'rb') as fp:
                    file_name = fp.read()
                cell.text = str(hex(zlib.crc32(file_name)))
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 5
                cell = table.cell(row, col)
                with open (i, 'rb') as fp:
                    file_name = fp.read()
                cell.text = str(hex(zlib.crc32(file_name)))
                row = row + 1

        #свойства документа
        for col in range (6):
            row = 0
            cell = table.cell(row, col)
            if col == 0:
                cell.text = str('Номер п/п')
            if col == 1:
                cell.text = str('Наименование электронного документа (файла)')
            if col == 2:
                cell.text = str('Дата и время последнего изменения документа (файла)')
            if col == 3:
                cell.text = str('Размер файла (байт)')
            if col == 4:
                cell.text = str('Версия')
            if col == 5:
                cell.text = str('CRC32')

        doct.add_paragraph()

        table2 = doct.add_table(rows = 2, cols = 3)
        table2.style = 'Table Grid'

        #создание ячеек
        for col in range(3):
            for row in range (2):
                cell = table2.cell(row, col)
                cell.text = str('')
        
        #свойства документа
        for col in range (3):
            row = 0
            cell = table2.cell(row, col)
            if col == 1:
                cell.text = str('Исполнитель')
            if col == 2:
                cell.text = str('Подпись')

        doct.add_paragraph()

        table3 = doct.add_table(rows = 2, cols = 2)
        table3.style = 'Table Grid'

        #создание ячеек
        for col in range(2):
            for row in range (2):
                cell = table3.cell(row, col)
                cell.text = str('')
        
        #сохранение искомой таблицы
        doct.save('CRC32 общая таблица.docx')
        
        lbl11.configure(text="Готово!2")
        lbl12.configure(text="Теперь готовый ИУЛ ждёт Вас в папке с программой")
    
    if selected.get() == 3:
        #print(folder_path.get())
        folder_path.get()
        paths = []
        folder = folder_path.get()
        #print(folder)
        i = 0
        m = 0
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                i = i + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                m = m + 1
        j = 0
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                j = j + 1
        k = 0
        n = 0
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                k = k + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                n = n + 1
        
        doct = Document()
        d = i + j + k + m + n
        #print(d)
        table = doct.add_table(rows = d + 1, cols = 5)
        table.style = 'Table Grid'

        #создание ячеек
        for col in range(5):
            for row in range (d + 1):
                cell = table.cell(row, col)
                cell.text = str(0)

        #нумерация документов с первого встреченного
        for row in range(1, d + 1):
            col = 0
            cell = table.cell(row, col)
            cell.text = str(row)

        #заполнение названий документов (1)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1

        #заполнение последних дат изменений (2)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                currentFile = docx.Document(currentFile)
                properties = currentFile.core_properties
                col = 2
                cell = table.cell(row, col)
                cell.text = str(properties.modified)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 2
                cell = table.cell(row, col)
                cell.text = str(0)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                pdf = pikepdf.Pdf.open(currentFile)
                docinfo = pdf.docinfo
                col = 2
                cell = table.cell(row, col)
                o = 0
                for key, value in docinfo.items():
                    if str(value).startswith("D:"):
                        value = transform_date(str(pdf.docinfo["/CreationDate"]))
                        if o == 0:
                            if value != '':
                                #print(str(value))
                                cell.text = str(value)
                            else:
                                cell.text = str('-')
                        if o == 2:
                            if value != '':
                                #print(str(value))
                                cell.text = str(value)
                            else:
                                cell.text = str('-')
                    o = o + 1
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                col = 2
                cell = table.cell(row, col)
                i = currentFile.resolve().stem + ".xlsx"
                #print(i)
                file = load_workbook(currentFile)
                #print(file.properties.modified)
                j = file.properties.modified
                #print(j)
                cell.text = str(j)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                col = 2
                cell = table.cell(row, col)
                i = currentFile.resolve().stem + ".xls"
                #print(i)
                assert olefile.isOleFile(currentFile)
                ole = olefile.OleFileIO(currentFile)
                meta = ole.get_metadata()
                #print('Дата последнего сохранения:  '+str(meta.last_saved_time))
                cell.text = str(meta.last_saved_time)
                row = row + 1
                ole.close()
        
        #заполнение размера документов (3)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1

        #заполнение версий документов (4)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                currentFile = docx.Document(currentFile)
                properties = currentFile.core_properties
                for col in range(5):
                    cell = table.cell(row, col)
                    if col == 4:
                        if properties.version != '':
                            cell.text = str(properties.version)
                        else:
                            cell.text = str('0')
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1

        #свойства документа
        for col in range (5):
            row = 0
            cell = table.cell(row, col)
            if col == 0:
                cell.text = str('Номер п/п')
            if col == 1:
                cell.text = str('Наименование электронного документа (файла)')
            if col == 2:
                cell.text = str('Дата и время последнего изменения документа (файла)')
            if col == 3:
                cell.text = str('Размер файла (байт)')
            if col == 4:
                cell.text = str('Версия')

        doct.add_paragraph()

        table1 = doct.add_table(rows = d + 1, cols = 2)
        table1.style = 'Table Grid'

        #создание ячеек
        for col in range(2):
            for row in range (d + 1):
                cell = table1.cell(row, col)
                cell.text = str(0)

        #нумерация документов с первого встреченного
        for row in range(1, d + 1):
            col = 0
            cell = table1.cell(row, col)
            cell.text = str(row)

        #заполнение контрольной суммы документов 
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 1
                cell = table1.cell(row, col)
                file_name = i
                cell.text = str(hashlib.md5(open(file_name, 'rb').read()).hexdigest())
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 1
                cell = table1.cell(row, col)
                file_name = i
                cell.text = str(hashlib.md5(open(file_name, 'rb').read()).hexdigest())
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 1
                cell = table1.cell(row, col)
                file_name = i
                cell.text = str(hashlib.md5(open(file_name, 'rb').read()).hexdigest())
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 1
                cell = table1.cell(row, col)
                file_name = i
                cell.text = str(hashlib.md5(open(file_name, 'rb').read()).hexdigest())
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 1
                cell = table1.cell(row, col)
                file_name = i
                cell.text = str(hashlib.md5(open(file_name, 'rb').read()).hexdigest())
                row = row + 1
        
        #свойства документа
        for col in range (2):
            row = 0
            cell = table1.cell(row, col)
            if col == 0:
                cell.text = str('Номер п/п')
            if col == 1:
                cell.text = str('MD5')

        doct.add_paragraph()

        table2 = doct.add_table(rows = 2, cols = 3)
        table2.style = 'Table Grid'

        #создание ячеек
        for col in range(3):
            for row in range (2):
                cell = table2.cell(row, col)
                cell.text = str('')
        
        #свойства документа
        for col in range (3):
            row = 0
            cell = table2.cell(row, col)
            if col == 1:
                cell.text = str('Исполнитель')
            if col == 2:
                cell.text = str('Подпись')

        doct.add_paragraph()

        table3 = doct.add_table(rows = 2, cols = 2)
        table3.style = 'Table Grid'

        #создание ячеек
        for col in range(2):
            for row in range (2):
                cell = table3.cell(row, col)
                cell.text = str('')
        
        #сохранение документа с таблицей
        doct.save('MD5 отдельной таблицей.docx')
        
        lbl11.configure(text="Готово!3")
        lbl12.configure(text="Теперь готовый ИУЛ ждёт Вас в папке с программой")
    
    if selected.get() == 4:
        #print(folder_path.get())
        folder_path.get()
        paths = []
        folder = folder_path.get()
        #print(folder)
        i = 0
        m = 0
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                i = i + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                m = m + 1
        j = 0
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                j = j + 1
        k = 0
        n = 0
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                k = k + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                #print(currentFile.resolve().stem)
                n = n + 1
        
        doct = Document()
        d = i + j + k + m + n
        #print(d)
        table = doct.add_table(rows = d + 1, cols = 5)
        table.style = 'Table Grid'

        #создание ячеек
        for col in range(5):
            for row in range (d + 1):
                cell = table.cell(row, col)
                cell.text = str(0)

        #нумерация документов с первого встреченного
        for row in range(1, d + 1):
            col = 0
            cell = table.cell(row, col)
            cell.text = str(row)

        #заполнение названий документов (1)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile.resolve().stem
                col = 1
                cell = table.cell(row, col)
                cell.text = str(i)
                row = row + 1

        #заполнение последних дат изменений (2)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                currentFile = docx.Document(currentFile)
                properties = currentFile.core_properties
                col = 2
                cell = table.cell(row, col)
                cell.text = str(properties.modified)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 2
                cell = table.cell(row, col)
                cell.text = str(0)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                pdf = pikepdf.Pdf.open(currentFile)
                docinfo = pdf.docinfo
                col = 2
                cell = table.cell(row, col)
                o = 0
                for key, value in docinfo.items():
                    if str(value).startswith("D:"):
                        value = transform_date(str(pdf.docinfo["/CreationDate"]))
                        if o == 0:
                            if value != '':
                                #print(str(value))
                                cell.text = str(value)
                            else:
                                cell.text = str('-')
                        if o == 2:
                            if value != '':
                                #print(str(value))
                                cell.text = str(value)
                            else:
                                cell.text = str('-')
                    o = o + 1
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                col = 2
                cell = table.cell(row, col)
                i = currentFile.resolve().stem + ".xlsx"
                #print(i)
                file = load_workbook(currentFile)
                #print(file.properties.modified)
                j = file.properties.modified
                #print(j)
                cell.text = str(j)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                col = 2
                cell = table.cell(row, col)
                i = currentFile.resolve().stem + ".xls"
                #print(i)
                assert olefile.isOleFile(currentFile)
                ole = olefile.OleFileIO(currentFile)
                meta = ole.get_metadata()
                #print('Дата последнего сохранения:  '+str(meta.last_saved_time))
                cell.text = str(meta.last_saved_time)
                row = row + 1
                ole.close()
        
        #заполнение размера документов (3)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 3
                cell = table.cell(row, col)
                statinfo = os.stat(i)
                cell.text = str(statinfo.st_size)
                row = row + 1

        #заполнение версий документов (4)
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                currentFile = docx.Document(currentFile)
                properties = currentFile.core_properties
                for col in range(5):
                    cell = table.cell(row, col)
                    if col == 4:
                        if properties.version != '':
                            cell.text = str(properties.version)
                        else:
                             cell.text = str('0')
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 4
                cell = table.cell(row, col)
                cell.text = str('0')
                row = row + 1

        #свойства документа
        for col in range (5):
            row = 0
            cell = table.cell(row, col)
            if col == 0:
                cell.text = str('Номер п/п')
            if col == 1:
                cell.text = str('Наименование электронного документа (файла)')
            if col == 2:
                cell.text = str('Дата и время последнего изменения документа (файла)')
            if col == 3:
                cell.text = str('Размер файла (байт)')
            if col == 4:
                cell.text = str('Версия')

        doct.add_paragraph()

        table1 = doct.add_table(rows = d + 1, cols = 2)
        table1.style = 'Table Grid'

        #создание ячеек
        for col in range(2):
            for row in range (d + 1):
                cell = table1.cell(row, col)
                cell.text = str(0)

        #нумерация документов с первого встреченного
        for row in range(1, d + 1):
            col = 0
            cell = table1.cell(row, col)
            cell.text = str(row)

        #заполнение контрольной суммы документов 
        row = 1
        if selected1.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.docx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 1
                cell = table1.cell(row, col)
                with open (i, 'rb') as fp:
                    file_name = fp.read()
                cell.text = str(hex(zlib.crc32(file_name)))
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.doc"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 1
                cell = table1.cell(row, col)
                with open (i, 'rb') as fp:
                    file_name = fp.read()
                cell.text = str(hex(zlib.crc32(file_name)))
                row = row + 1
        if selected2.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.pdf"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 1
                cell = table1.cell(row, col)
                with open (i, 'rb') as fp:
                    file_name = fp.read()
                cell.text = str(hex(zlib.crc32(file_name)))
                row = row + 1
        if selected3.get() == 1:
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xlsx"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 1
                cell = table1.cell(row, col)
                with open (i, 'rb') as fp:
                    file_name = fp.read()
                cell.text = str(hex(zlib.crc32(file_name)))
                row = row + 1
            currentDirectory = pathlib.Path(folder)
            currentPattern = "*.xls"
            for currentFile in currentDirectory.glob(currentPattern):
                i = currentFile
                col = 1
                cell = table1.cell(row, col)
                with open (i, 'rb') as fp:
                    file_name = fp.read()
                cell.text = str(hex(zlib.crc32(file_name)))
                row = row + 1
        
        #свойства документа
        for col in range (2):
            row = 0
            cell = table1.cell(row, col)
            if col == 0:
                cell.text = str('Номер п/п')
            if col == 1:
                cell.text = str('CRC32')

        doct.add_paragraph()

        table2 = doct.add_table(rows = 2, cols = 3)
        table2.style = 'Table Grid'

        #создание ячеек
        for col in range(3):
            for row in range (2):
                cell = table2.cell(row, col)
                cell.text = str('')
        
        #свойства документа
        for col in range (3):
            row = 0
            cell = table2.cell(row, col)
            if col == 1:
                cell.text = str('Исполнитель')
            if col == 2:
                cell.text = str('Подпись')

        doct.add_paragraph()

        table3 = doct.add_table(rows = 2, cols = 2)
        table3.style = 'Table Grid'

        #создание ячеек
        for col in range(2):
            for row in range (2):
                cell = table3.cell(row, col)
                cell.text = str('')
        
        #сохранение документа с таблицей
        doct.save('CRC32 отдельной таблицей.docx')

        lbl11.configure(text="Готово!4")
        lbl12.configure(text="Теперь готовый ИУЛ ждёт Вас в папке с программой")

def browse_button():
    filename = filedialog.askdirectory()
    global folder_path
    if (filename != ''):
        folder_path.set(filename)
        lbl1.configure(text=filename)
        lbl.configure(text="Выберите нужный вариант и нажмите сюда --->")
    else:
        lbl1.configure(text="Папка не выбрана!")
    #print(filename)

def f1():
    #lbl1.configure(text="Папка выбрана!")
    window0 = tk.Tk()
    window0.title("Выбор папки")
    window0.geometry('400x250')
    folder_path = StringVar()
    lbl0 = Label(master=window0,textvariable=folder_path)
    lbl0.grid(row=5, column=1)
    lbl = Label(window0)
    lbl.grid(column=1, row=0)
    button2 = Button(window0, text="Выбрать папку", command=browse_button)
    button2.grid(row=0, column=1)
    window0.mainloop()

window = tk.Tk()
window.title("Генерация ИУЛ")
window.geometry('800x300')
folder_path = StringVar()
selected = IntVar()
selected1 = IntVar()
selected2 = IntVar()
selected3 = IntVar()
lbl7 = Label(window)
lbl1 = Label(window)
lbl9 = Label(window)
lbl2 = Label(window)
lbl8 = Label(window)
lbl7.grid(column=1, row=0)
lbl1.grid(column=1, row=1)
lbl9.grid(column=1, row=2)
lbl2.grid(column=1, row=7)
lbl8.grid(column=1, row=4)
lbl7.configure(text="Сначала выберите папку")
lbl9.configure(text="Затем выберете тип обрабатываемых файлов")
lbl2.configure(text="Пока не выбрали папку не генерировать ИУЛ!!!")
lbl8.configure(text="Затем выберите вид таблицы")
btn1 = Button(window, text="Выбрать папку", command=f1)
btn1.grid(column=2, row=0)
rad1 = Radiobutton(window, text='MD5 общая таблица', value=1, variable=selected)  
rad2 = Radiobutton(window, text='CRC32 общая таблица', value=2, variable=selected)  
rad3 = Radiobutton(window, text='MD5 отдельной таблицей', value=3, variable=selected)
rad4 = Radiobutton(window, text='CRC32 отдельной таблицей', value=4, variable=selected)

rad5 = Checkbutton(window, text='DOCX (и DOC)', variable=selected1)
rad6 = Checkbutton(window, text='PDF', variable=selected2)
rad7 = Checkbutton(window, text='XLSX (и XLS)', variable=selected3)
btn = Button(window, text="Сгенерировать ИУЛ", command = clicked)
lbl = Label(window)
lbl11 = Label(window)
lbl12 = Label(window)
rad1.grid(column=1, row=5)
rad2.grid(column=2, row=5)
rad3.grid(column=1, row=6)
rad4.grid(column=2, row=6)

rad5.grid(column=1, row=3)
rad6.grid(column=2, row=3)
rad7.grid(column=3, row=3)

btn.grid(column=2, row=7)
lbl.grid(column=1, row=7)
lbl11.grid(column=1,row=8)
lbl12.grid(column=2,row=8)

window.mainloop()

#проверка, если нужно
#print('Готово!')
