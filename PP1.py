import tkinter as tk
import sys
import pathlib
import os
import docx
import docx2pdf

from tkinter import messagebox as mb
from tkinter import filedialog
from tkinter import *
from os import path
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader

#sys.path.append('../../')

def f_count0(text, sequence):
    count = 0
    for i in range(len(text) - len(sequence) + 1):
        if text[i:i+len(sequence)] == sequence:
            count += 1
    return count

def f_count1(file_path, sequence):
    count = 0
    o = ""
    with open(file_path, "r") as file:
        lines = file.readlines()
        for j, line in enumerate(lines):
            line = line.rstrip("\n")
            for i in range(len(line) - len(sequence) + 1):
                if line[i:i+len(sequence)] == sequence:
                    count += 1
                    o += f"Последовательность '{sequence}' найдена в строчке {j + 1}" + "\n"
    lbl14_0.configure(text= o)
    return count

def f_count2(file_path, sequence):
    count = 0
    o = ""
    doc = Document(file_path)
    paragraphs = doc.paragraphs
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text
        if sequence in text:
            count += 1
            o += f"Последовательность '{sequence}' найдена в параграфе {i + 1}" + "\n"
    lbl14_0.configure(text= o)
    return count

def fbtn1():
    if selected1.get() == 1:
        lbl3.configure(text="Выбор сделан!")
        btn2 = Button(window, text="Выбрать файл", command = fbtn2)
        btn2.grid(column=1, row=5)
        rad2.grid_remove()
    
    if selected1.get() == 2:
        lbl3.configure(text="Выбор сделан!")
        lbl4_1.configure(text="Вставьте или введите свой текст: ")
        entry0.grid(column=2, row=6)
        lbl13.configure(text="Введите нужную Вам фразу или слово: ")
        entry00.grid(column=2, row=7)
        rad1.grid_remove()
        btn5 = Button(window, text="Выполнить проверку", command=fbtn5)
        btn5.grid(column=2, row=8)

def fbtn2():
    window0 = tk.Tk()
    window0.title("Выбор файла")
    window0.geometry('300x200')
    folder_path = StringVar()
    lbl0 = Label(master=window0,textvariable=folder_path)
    lbl0.grid(column=1, row=5)
    lbl = Label(window0)
    lbl.grid(column=1, row=0)
    btn2 = Button(window0, text="Выбрать файл", command=browse_button)
    btn2.grid(column=1, row=0)
    window0.mainloop()

def fbtn4_0():
    btn4 = Button(window, text="Выполнить проверку", command=fbtn4_0)
    btn4.grid(column=1, row=12)
    global folder_path
    folder_path.get()
    paths = []
    folder = folder_path.get()
    t = entry.get()
    lbl9['text'] = t
    with open (fn(folder), 'r') as file:
        lbl11.configure(text= "Количество вхождений последовательности в тексте: " + str(f_count1(fn(folder_path.get()), t)))
    
def fbtn4_1():
    btn4 = Button(window, text="Выполнить проверку", command=fbtn4_1)
    btn4.grid(column=1, row=12)
    global folder_path
    folder_path.get()
    paths = []
    folder = folder_path.get()
    t = entry.get()
    lbl9['text'] = t
    doc = docx.Document(fn(folder))
    ftext = []
    for para in doc.paragraphs:
        ftext.append(para.text)
    lbl11.configure(text= "Количество вхождений последовательности в тексте: " + str(f_count2(fn(folder_path.get()), t)))

def fbtn5():
    s = entry00.get()
    t = entry0.get()
    lbl11.configure(text= "Количество вхождений последовательности в тексте: " + str(f_count0(t, s)))

def fn(filename):
    parts = filename.split("'")
    if len(parts) > 1:
        return parts[1]
    else:
        return ''

def browse_button():
    filename = filedialog.askopenfile()
    global folder_path
    if (filename != ''):
        lbl4_0.configure(text="Расположение файла: ")

        folder_path.set(filename)
        folder = folder_path.get()
        lbl5.configure(text=fn(folder))
        
        lbl0.configure(text="Файл выбран")

        entry.grid(column=2, row=11)
        lbl8.configure(text="Введите нужную Вам фразу или слово: ")

        f, s = os.path.splitext(fn(folder))

        if s == ".txt":
            btn4 = Button(window, text="Выполнить проверку", command=fbtn4_0)
            btn4.grid(column=1, row=12)
        if s == ".docx":
            btn4 = Button(window, text="Выполнить проверку", command=fbtn4_1)
            btn4.grid(column=1, row=12)
    else:
        lbl0.configure(text="Файл не выбран!")

window = tk.Tk()
window.title("TextCheck")
window.geometry('800x400')

folder_path = StringVar()
s = StringVar()
selected1 = IntVar()
selected2 = IntVar()

entry = Entry()
entry0 = Entry()
entry00 = Entry()

lbl1 = Label(window)
lbl2 = Label(window)
lbl1.configure(text="Эпичное название программы")
lbl2.configure(text="Вам нужно проверить файл с текстом\n или у Вас есть текст не в файле, который нужно проверить?")
lbl1.grid(column=1, row=1)
lbl2.grid(column=1, row=2)

rad1 = Radiobutton(window, text='Файл с текстом', value=1, variable=selected1)
rad2 = Radiobutton(window, text='Текст не в файле', value=2, variable=selected1)
rad1.grid(column=1, row=3)
rad2.grid(column=2, row=3)

btn1 = Button(window, text="Выбрать", command = fbtn1)
btn1.grid(column=1, row=4)

lbl3 = Label(window)
lbl3.grid(column=2, row=4)

lbl4_0 = Label(window)
lbl4_1 = Label(window)
lbl4_0.grid(column=1, row=9)
lbl4_1.grid(column=1, row=6)

lbl0 = Label(window)
lbl0.grid(column=2, row=5)

lbl5 = Label(window)
lbl6 = Label(window)
lbl5.grid(column=1, row=10)
lbl6.grid(column=1, row=8)

lbl7 = Label(window)
lbl7.grid(column=2, row=10)

lbl8 = Label(window)
lbl8.grid(column=1, row=11)

lbl9 = Label(window)
lbl9.grid(column=2, row=12)

lbl10 = Label(window)
lbl10.grid(column=1, row=13)

lbl11 = Label(window)
lbl11.grid(column=2, row=13)

lbl12 = Label(window)
lbl12.grid(column=2, row=9)

lbl13 = Label(window)
lbl13.grid(column=1, row=7)

lbl14_0 = Label(window)
lbl14_0.grid(column=2, row=14)

lbl14_1 = Label(window)
lbl14_1.grid(column=1, row=14)

window.mainloop()

#проверка, если нужно
#print('Готово!')
